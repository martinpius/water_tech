from __future__ import annotations

import math
import re
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE_FILE = ROOT / "original_data.xlsx"
OUTPUT_DIR = ROOT / "maturity_quantity_quality_outputs"
PLOTS_DIR = OUTPUT_DIR / "plots"
REPORT_FILE = OUTPUT_DIR / "report.txt"
DOCX_FILE = ROOT / "Maturity, quantity, and integration interpretation.docx"

NUMERIC_COLUMNS = [
    "Awareness Score",
    "Number of Tools Adopted",
    "Integration Score",
    "Performance Score",
    "Maturity Score",
]


def parse_assigned_numeric(value: object) -> float | None:
    if pd.isna(value):
        return None
    if isinstance(value, (int, float, np.integer, np.floating)):
        return float(value)
    text = str(value).strip()
    match = re.search(r"Assigned\s*([0-9.]+)", text)
    if match:
        return float(match.group(1))
    try:
        return float(text)
    except ValueError:
        return None


def normalize_tool(tool: str) -> str:
    text = str(tool).strip().strip(".")
    replacements = {
        "AMI/Smart meters": "AMI",
        "IoT": "IoT sensors",
        "Sensors": "IoT sensors",
        "Automated devices": "IoT sensors",
        "Data analytics": "Analytics",
        "Billing systems": "Billing",
        "Distribution systems": "Distribution",
        "Quality monitoring": "Quality",
    }
    return replacements.get(text, text)


def split_tools(text: object) -> list[str]:
    if pd.isna(text):
        return []
    parts = re.split(r";|,", str(text))
    tools = [normalize_tool(part) for part in parts if str(part).strip()]
    return [tool for tool in tools if tool]


def load_data() -> pd.DataFrame:
    raw = pd.read_excel(SOURCE_FILE, header=None)
    frame = raw.iloc[3:].copy().reset_index(drop=True)
    frame.columns = raw.iloc[2]
    frame = frame.loc[:, [column for column in frame.columns if not pd.isna(column)]]
    frame.columns = [str(column).strip() for column in frame.columns]
    for column in frame.columns:
        frame[column] = frame[column].ffill()
    for column in NUMERIC_COLUMNS:
        frame[column] = frame[column].map(parse_assigned_numeric)
    frame["row_id"] = np.arange(1, len(frame) + 1)
    frame["is_adopted"] = frame["Adoption Status"].astype(str).str.strip().eq("Adopted")
    frame["integration_pure_full"] = frame["System Integration Level"].astype(str).str.strip().eq("Full")
    frame["tool_list"] = frame["Digital Tools Adopted"].map(split_tools)
    frame["tool_count_listed"] = frame["tool_list"].map(len)
    return frame


def correlation_table(df: pd.DataFrame) -> pd.DataFrame:
    cols = ["Number of Tools Adopted", "Integration Score", "Maturity Score", "Performance Score"]
    return df[cols].corr().round(3)


def contradiction_tables(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    adopted = df[df["is_adopted"]].copy()
    low_full = adopted[(adopted["Maturity Score"] <= 1) & (adopted["integration_pure_full"])].copy()
    high_not_full = adopted[(adopted["Maturity Score"] >= 3) & (~adopted["integration_pure_full"])].copy()
    keep = [
        "Institution",
        "Digital Tools Adopted",
        "Number of Tools Adopted",
        "System Integration Level",
        "Integration Score",
        "Technological Maturity",
        "Maturity Score",
        "Performance Score",
    ]
    return low_full[keep], high_not_full[keep]


def maturity_integration_summary(df: pd.DataFrame) -> pd.DataFrame:
    adopted = df[df["is_adopted"]].copy()
    summary = (
        adopted.groupby(["Maturity Score", "System Integration Level"], dropna=False)
        .agg(
            institutions=("Institution", "size"),
            mean_quantity=("Number of Tools Adopted", "mean"),
            mean_performance=("Performance Score", "mean"),
        )
        .reset_index()
        .sort_values(["Maturity Score", "System Integration Level"])
    )
    summary["mean_quantity"] = summary["mean_quantity"].round(2)
    summary["mean_performance"] = summary["mean_performance"].round(2)
    return summary


def performance_summary(df: pd.DataFrame) -> pd.DataFrame:
    adopted = df[df["is_adopted"]].copy()
    rows = []
    for maturity, sub in adopted.groupby("Maturity Score"):
        rows.append(
            {
                "maturity_score": maturity,
                "n": len(sub),
                "mean_quantity": round(float(sub["Number of Tools Adopted"].mean()), 2),
                "mean_integration": round(float(sub["Integration Score"].mean()), 2),
                "mean_performance": round(float(sub["Performance Score"].mean()), 2),
            }
        )
    return pd.DataFrame(rows).sort_values("maturity_score")


def tool_level_summary(df: pd.DataFrame) -> pd.DataFrame:
    adopted = df[df["is_adopted"]].copy()
    long_rows: list[dict[str, object]] = []
    for _, row in adopted.iterrows():
        for tool in row["tool_list"]:
            long_rows.append(
                {
                    "tool": tool,
                    "institution": row["Institution"],
                    "quantity": row["Number of Tools Adopted"],
                    "integration_score": row["Integration Score"],
                    "maturity_score": row["Maturity Score"],
                    "performance_score": row["Performance Score"],
                }
            )
    long_df = pd.DataFrame(long_rows)
    summary = (
        long_df.groupby("tool")
        .agg(
            mentions=("tool", "size"),
            mean_quantity=("quantity", "mean"),
            mean_integration=("integration_score", "mean"),
            mean_maturity=("maturity_score", "mean"),
            mean_performance=("performance_score", "mean"),
        )
        .reset_index()
        .sort_values(["mentions", "mean_maturity"], ascending=[False, False])
    )
    for column in ["mean_quantity", "mean_integration", "mean_maturity", "mean_performance"]:
        summary[column] = summary[column].round(2)
    return summary


def save_svg(path: Path, parts: list[str]) -> None:
    parts.append("</svg>\n")
    path.write_text("\n".join(parts), encoding="utf-8")


def svg_header(width: int, height: int, title: str | None = None, subtitle: str | None = None) -> list[str]:
    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
        '<rect width="100%" height="100%" fill="#fcfbf7"/>',
    ]
    if title:
        parts.append(f'<text x="{width/2}" y="28" text-anchor="middle" font-size="21" font-family="Georgia">{title}</text>')
    if subtitle:
        parts.append(
            f'<text x="{width/2}" y="48" text-anchor="middle" font-size="11" font-family="Arial" fill="#666">{subtitle}</text>'
        )
    return parts


def safe_label(text: object) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def heatmap_svg(summary: pd.DataFrame, path: Path) -> None:
    width, height = 980, 560
    margin_left, margin_top = 180, 95
    cell_w, cell_h = 145, 78
    parts = svg_header(
        width,
        height,
        "Figure 1. Maturity by integration status",
        "Cells show institution count and mean performance among adopted cases",
    )
    maturity_levels = sorted(summary["Maturity Score"].dropna().unique())
    integration_levels = [
        "Full",
        "Partial → Full",
        "Partial",
        "None → Full",
        "None → Partial",
    ]
    max_n = max(int(summary["institutions"].max()), 1)
    for idx, level in enumerate(integration_levels):
        x = margin_left + idx * cell_w + cell_w / 2
        parts.append(
            f'<text x="{x:.1f}" y="{margin_top-16}" text-anchor="middle" font-size="12" font-family="Arial">{safe_label(level)}</text>'
        )
    for ridx, maturity in enumerate(maturity_levels):
        y = margin_top + ridx * cell_h + cell_h / 2 + 5
        parts.append(
            f'<text x="{margin_left-14}" y="{y:.1f}" text-anchor="end" font-size="12" font-family="Arial">Maturity {maturity:.0f}</text>'
        )
        for cidx, level in enumerate(integration_levels):
            x0 = margin_left + cidx * cell_w
            y0 = margin_top + ridx * cell_h
            row = summary[
                (summary["Maturity Score"] == maturity) & (summary["System Integration Level"] == level)
            ]
            if row.empty:
                count = 0
                perf = None
            else:
                count = int(row.iloc[0]["institutions"])
                perf = float(row.iloc[0]["mean_performance"])
            intensity = count / max_n
            fill = f"rgb({245 - int(70 * intensity)},{241 - int(40 * intensity)},{230 - int(120 * intensity)})"
            parts.append(f'<rect x="{x0}" y="{y0}" width="{cell_w-10}" height="{cell_h-10}" rx="8" fill="{fill}" stroke="#d7d1c4"/>')
            parts.append(
                f'<text x="{x0 + (cell_w-10)/2:.1f}" y="{y0+30:.1f}" text-anchor="middle" font-size="16" font-family="Georgia">{count}</text>'
            )
            perf_text = "NA" if perf is None else f"Perf {perf:.2f}"
            parts.append(
                f'<text x="{x0 + (cell_w-10)/2:.1f}" y="{y0+53:.1f}" text-anchor="middle" font-size="11" font-family="Arial" fill="#444">{perf_text}</text>'
            )
    legend_x, legend_y = width - 220, height - 65
    for i in range(5):
        intensity = i / 4
        fill = f"rgb({245 - int(70 * intensity)},{241 - int(40 * intensity)},{230 - int(120 * intensity)})"
        parts.append(f'<rect x="{legend_x + i*28}" y="{legend_y}" width="24" height="16" fill="{fill}" stroke="#d7d1c4"/>')
    parts.append(f'<text x="{legend_x}" y="{legend_y-8}" font-size="11" font-family="Arial">Higher shading = more institutions</text>')
    save_svg(path, parts)


def scatter_svg(df: pd.DataFrame, path: Path) -> None:
    adopted = df[df["is_adopted"]].copy()
    width, height = 920, 590
    margin = 80
    plot_bottom = 470
    parts = svg_header(width, height, None)
    x = adopted["Number of Tools Adopted"].to_numpy(float)
    y = adopted["Performance Score"].to_numpy(float)
    x_lo, x_hi = float(x.min()), float(x.max())
    y_lo, y_hi = float(y.min()), float(y.max())

    def x_map(v: float) -> float:
        return margin + (v - x_lo) * (width - 2 * margin) / max(1e-9, x_hi - x_lo)

    def y_map(v: float) -> float:
        return plot_bottom - (v - y_lo) * (plot_bottom - margin) / max(1e-9, y_hi - y_lo)

    parts.extend(
        [
            f'<line x1="{margin}" y1="{plot_bottom}" x2="{width-margin}" y2="{plot_bottom}" stroke="#222"/>',
            f'<line x1="{margin}" y1="{margin}" x2="{margin}" y2="{plot_bottom}" stroke="#222"/>',
        ]
    )
    for tick in range(int(x_lo), int(x_hi) + 1):
        tx = x_map(float(tick))
        parts.append(f'<line x1="{tx:.1f}" y1="{margin}" x2="{tx:.1f}" y2="{plot_bottom}" stroke="#e2ddd3" stroke-width="1"/>')
    for tick in np.arange(round(y_lo, 1), y_hi + 0.01, 0.4):
        ty = y_map(float(tick))
        parts.append(f'<line x1="{margin}" y1="{ty:.1f}" x2="{width-margin}" y2="{ty:.1f}" stroke="#e2ddd3" stroke-width="1"/>')
    colors = {0.0: "#d95f02", 1.0: "#f2b447", 2.0: "#4e79a7", 3.0: "#2a9d8f", 4.0: "#5b2c83"}
    contradiction_names = {
        "Internal Drainage Basin Water Board (IDBWB)": "IDBWB",
        "Lake Tanganyika Basin Water Board (LTBWB)": "LTBWB",
        "TUWASA ": "TUWASA",
        "Bugando Medical Centre ": "Bugando",
        "DAWASA ": "DAWASA",
        "DUWASA": "DUWASA",
        "MWAUWASA": "MWAUWASA",
        "Davis and Shirtliff Ltd ": "Davis & Shirtliff",
    }
    for _, row in adopted.iterrows():
        cx = x_map(float(row["Number of Tools Adopted"]))
        cy = y_map(float(row["Performance Score"]))
        maturity = float(row["Maturity Score"])
        radius = 6 + 2.5 * float(row["Integration Score"])
        stroke = "#111" if row["integration_pure_full"] else "#8b0000"
        parts.append(
            f'<circle cx="{cx:.1f}" cy="{cy:.1f}" r="{radius:.1f}" fill="{colors.get(maturity, "#888")}" fill-opacity="0.78" stroke="{stroke}" stroke-width="1.4"/>'
        )
        label = contradiction_names.get(str(row["Institution"]))
        if label:
            parts.append(
                f'<text x="{cx+8:.1f}" y="{cy-9:.1f}" font-size="10.5" font-family="Arial" fill="#222">{safe_label(label)}</text>'
            )
    slope, intercept = np.polyfit(x, y, 1)
    x_line = np.array([x_lo, x_hi])
    y_line = intercept + slope * x_line
    parts.append(
        f'<line x1="{x_map(float(x_line[0])):.1f}" y1="{y_map(float(y_line[0])):.1f}" x2="{x_map(float(x_line[1])):.1f}" y2="{y_map(float(y_line[1])):.1f}" stroke="#333" stroke-width="2"/>'
    )
    for tick in range(int(x_lo), int(x_hi) + 1):
        tx = x_map(float(tick))
        parts.append(f'<line x1="{tx:.1f}" y1="{plot_bottom}" x2="{tx:.1f}" y2="{plot_bottom+6}" stroke="#222"/>')
        parts.append(f'<text x="{tx:.1f}" y="{plot_bottom+22}" text-anchor="middle" font-size="11">{tick}</text>')
    for tick in np.arange(round(y_lo, 1), y_hi + 0.01, 0.4):
        ty = y_map(float(tick))
        parts.append(f'<line x1="{margin-6}" y1="{ty:.1f}" x2="{margin}" y2="{ty:.1f}" stroke="#222"/>')
        parts.append(f'<text x="{margin-10}" y="{ty+4:.1f}" text-anchor="end" font-size="11">{tick:.1f}</text>')
    parts.append(f'<text x="{width/2}" y="{plot_bottom+52}" text-anchor="middle" font-size="13" font-family="Arial">Number of tools adopted (quantity)</text>')
    parts.append(f'<text x="24" y="{(plot_bottom+margin)/2}" transform="rotate(-90 24 {(plot_bottom+margin)/2})" text-anchor="middle" font-size="13" font-family="Arial">Performance score</text>')
    legend_x, legend_y = width - 210, plot_bottom - 120
    parts.append(f'<rect x="{legend_x}" y="{legend_y}" width="160" height="118" rx="10" fill="#fffefb" stroke="#d7d1c4"/>')
    parts.append(f'<text x="{legend_x+12}" y="{legend_y+18}" font-size="11" font-family="Arial">Maturity score</text>')
    for idx, level in enumerate([0.0, 1.0, 2.0, 3.0, 4.0]):
        cy = legend_y + 36 + idx * 18
        parts.append(f'<circle cx="{legend_x+18}" cy="{cy}" r="6" fill="{colors[level]}" stroke="#666"/>')
        parts.append(f'<text x="{legend_x+32}" y="{cy+4}" font-size="11" font-family="Arial">{int(level)}</text>')
    parts.append(f'<text x="{legend_x+12}" y="{legend_y+106}" font-size="10.5" font-family="Arial">Dark outline = not pure full</text>')
    save_svg(path, parts)


def tool_map_svg(tool_summary: pd.DataFrame, path: Path) -> None:
    top = tool_summary[tool_summary["mentions"] >= 2].copy().sort_values("mentions", ascending=False).head(12)
    width, height = 940, 560
    margin = 85
    parts = svg_header(width, height, None)
    x = top["mean_maturity"].to_numpy(float)
    y = top["mean_integration"].to_numpy(float)
    x_lo, x_hi = max(0.0, float(x.min()) - 0.2), min(4.2, float(x.max()) + 0.2)
    y_lo, y_hi = max(0.0, float(y.min()) - 0.15), min(2.1, float(y.max()) + 0.15)

    def x_map(v: float) -> float:
        return margin + (v - x_lo) * (width - 2 * margin) / max(1e-9, x_hi - x_lo)

    def y_map(v: float) -> float:
        return height - margin - (v - y_lo) * (height - 2 * margin) / max(1e-9, y_hi - y_lo)

    parts.extend(
        [
            f'<line x1="{margin}" y1="{height-margin}" x2="{width-margin}" y2="{height-margin}" stroke="#222"/>',
            f'<line x1="{margin}" y1="{margin}" x2="{margin}" y2="{height-margin}" stroke="#222"/>',
        ]
    )
    x_grid_ticks = np.arange(math.ceil(x_lo * 2) / 2, math.floor(x_hi * 2) / 2 + 0.001, 0.5)
    y_grid_ticks = np.arange(math.ceil(y_lo * 4) / 4, math.floor(y_hi * 4) / 4 + 0.001, 0.25)
    for tick in x_grid_ticks:
        tx = x_map(float(tick))
        parts.append(f'<line x1="{tx:.1f}" y1="{margin}" x2="{tx:.1f}" y2="{height-margin}" stroke="#e2ddd3" stroke-width="1"/>')
    for tick in y_grid_ticks:
        ty = y_map(float(tick))
        parts.append(f'<line x1="{margin}" y1="{ty:.1f}" x2="{width-margin}" y2="{ty:.1f}" stroke="#e2ddd3" stroke-width="1"/>')
    for _, row in top.iterrows():
        cx = x_map(float(row["mean_maturity"]))
        cy = y_map(float(row["mean_integration"]))
        radius = 7 + 1.6 * math.sqrt(float(row["mentions"]))
        parts.append(
            f'<circle cx="{cx:.1f}" cy="{cy:.1f}" r="{radius:.1f}" fill="#4e79a7" fill-opacity="0.7" stroke="#1f3a5f" stroke-width="1.2"/>'
        )
        parts.append(
            f'<text x="{cx+10:.1f}" y="{cy-8:.1f}" font-size="11" font-family="Arial">{safe_label(row["tool"])}</text>'
        )
    x_major_ticks = np.arange(math.ceil(x_lo), math.floor(x_hi) + 0.001, 1.0)
    y_major_ticks = np.arange(math.ceil(y_lo * 2) / 2, math.floor(y_hi * 2) / 2 + 0.001, 0.5)
    for tick in x_major_ticks:
        tx = x_map(float(tick))
        parts.append(f'<line x1="{tx:.1f}" y1="{height-margin}" x2="{tx:.1f}" y2="{height-margin+6}" stroke="#222"/>')
        parts.append(f'<text x="{tx:.1f}" y="{height-margin+20}" text-anchor="middle" font-size="11">{tick:.0f}</text>')
    for tick in y_major_ticks:
        ty = y_map(float(tick))
        parts.append(f'<line x1="{margin-6}" y1="{ty:.1f}" x2="{margin}" y2="{ty:.1f}" stroke="#222"/>')
        parts.append(f'<text x="{margin-10}" y="{ty+4:.1f}" text-anchor="end" font-size="11">{tick:.1f}</text>')
    parts.append(f'<text x="{width/2}" y="{height-18}" text-anchor="middle" font-size="13" font-family="Arial">Average maturity score (quality)</text>')
    parts.append(f'<text x="24" y="{height/2}" transform="rotate(-90 24 {height/2})" text-anchor="middle" font-size="13" font-family="Arial">Average integration score</text>')
    save_svg(path, parts)


def contradiction_panel_svg(low_full: pd.DataFrame, high_not_full: pd.DataFrame, path: Path) -> None:
    width, height = 1040, 620
    parts = svg_header(width, height, None)
    x0, y0 = 60, 52
    col0, col1, col2 = 280, 330, 330
    row_h = 46
    header_h = 54
    rows = [
        ("Cases", str(len(low_full)), str(len(high_not_full))),
        ("Mean quantity", f"{low_full['Number of Tools Adopted'].mean():.1f}", f"{high_not_full['Number of Tools Adopted'].mean():.1f}"),
        ("Mean integration", f"{low_full['Integration Score'].mean():.1f}", f"{high_not_full['Integration Score'].mean():.1f}"),
        ("Mean performance", f"{low_full['Performance Score'].mean():.1f}", f"{high_not_full['Performance Score'].mean():.1f}"),
        ("Dominant tools", "", ""),
    ]

    def tool_counts(frame: pd.DataFrame) -> list[tuple[str, int]]:
        counts: dict[str, int] = {}
        for text in frame["Digital Tools Adopted"]:
            for tool in split_tools(text):
                counts[tool] = counts.get(tool, 0) + 1
        return sorted(counts.items(), key=lambda item: (-item[1], item[0]))

    left_tools = tool_counts(low_full)[:8]
    right_tools = tool_counts(high_not_full)[:8]

    parts.append(f'<rect x="{x0}" y="{y0}" width="{col0+col1+col2}" height="{header_h + row_h*len(rows) + 320}" rx="18" fill="#fffdfa" stroke="#d8d2c8"/>')
    parts.append(f'<rect x="{x0}" y="{y0}" width="{col0}" height="{header_h}" rx="18" fill="#f3efe6" stroke="#d8d2c8"/>')
    parts.append(f'<rect x="{x0+col0}" y="{y0}" width="{col1}" height="{header_h}" fill="#fbf4ea" stroke="#d8c2a6"/>')
    parts.append(f'<rect x="{x0+col0+col1}" y="{y0}" width="{col2}" height="{header_h}" rx="18" fill="#edf7f4" stroke="#b9d9cd"/>')
    parts.append(f'<text x="{x0+24}" y="{y0+34}" font-size="15" font-family="Georgia">Comparison dimension</text>')
    parts.append(f'<text x="{x0+col0+col1/2}" y="{y0+22}" text-anchor="middle" font-size="15" font-family="Georgia">Low maturity + full integration</text>')
    parts.append(f'<text x="{x0+col0+col1/2}" y="{y0+40}" text-anchor="middle" font-size="11" font-family="Arial" fill="#6b5744">Simple tool bundles, high claimed integration</text>')
    parts.append(f'<text x="{x0+col0+col1+col2/2}" y="{y0+22}" text-anchor="middle" font-size="15" font-family="Georgia">High maturity + not pure full integration</text>')
    parts.append(f'<text x="{x0+col0+col1+col2/2}" y="{y0+40}" text-anchor="middle" font-size="11" font-family="Arial" fill="#45685f">Complex tool portfolios, transitional integration</text>')

    for idx, (label, left_val, right_val) in enumerate(rows):
        ry = y0 + header_h + idx * row_h
        fill = "#fcfaf5" if idx % 2 == 0 else "#fffdfa"
        parts.append(f'<rect x="{x0}" y="{ry}" width="{col0+col1+col2}" height="{row_h}" fill="{fill}" stroke="#ebe4d8"/>')
        parts.append(f'<line x1="{x0+col0}" y1="{ry}" x2="{x0+col0}" y2="{ry+row_h}" stroke="#e3d7c3"/>')
        parts.append(f'<line x1="{x0+col0+col1}" y1="{ry}" x2="{x0+col0+col1}" y2="{ry+row_h}" stroke="#d5e6df"/>')
        parts.append(f'<text x="{x0+18}" y="{ry+29}" font-size="12.5" font-family="Arial">{safe_label(label)}</text>')
        if left_val:
            parts.append(f'<text x="{x0+col0+24}" y="{ry+29}" font-size="16" font-family="Georgia">{safe_label(left_val)}</text>')
        if right_val:
            parts.append(f'<text x="{x0+col0+col1+24}" y="{ry+29}" font-size="16" font-family="Georgia">{safe_label(right_val)}</text>')

    cloud_top = y0 + header_h + row_h * len(rows) + 24
    cloud_h = 250
    parts.append(f'<rect x="{x0+col0+18}" y="{cloud_top}" width="{col1-36}" height="{cloud_h}" rx="16" fill="#fff7ec" stroke="#e7d3b9"/>')
    parts.append(f'<rect x="{x0+col0+col1+18}" y="{cloud_top}" width="{col2-36}" height="{cloud_h}" rx="16" fill="#f2fbf7" stroke="#cfe5dd"/>')

    def draw_word_cluster(items: list[tuple[str, int]], panel_x: int, panel_w: int, palette: list[str]) -> None:
        positions = [
            (0.50, 0.22), (0.30, 0.40), (0.70, 0.38), (0.52, 0.56),
            (0.25, 0.64), (0.74, 0.63), (0.36, 0.80), (0.64, 0.79),
        ]
        max_count = max((count for _, count in items), default=1)
        cx = panel_x + panel_w / 2
        cy = cloud_top + cloud_h / 2
        for radius, opacity in [(105, 0.18), (72, 0.28), (42, 0.42)]:
            parts.append(f'<circle cx="{cx:.1f}" cy="{cy:.1f}" r="{radius}" fill="#ffffff" fill-opacity="{opacity}" stroke="none"/>')
        for idx, (tool, count) in enumerate(items):
            px = panel_x + positions[idx][0] * panel_w
            py = cloud_top + positions[idx][1] * cloud_h
            size = 15 + int(16 * count / max_count)
            rotate = -8 if idx % 3 == 0 else (8 if idx % 4 == 0 else 0)
            color = palette[idx % len(palette)]
            parts.append(
                f'<text x="{px:.1f}" y="{py:.1f}" text-anchor="middle" font-size="{size}" font-family="Arial" font-weight="600" fill="{color}" transform="rotate({rotate} {px:.1f} {py:.1f})">{safe_label(tool)}</text>'
            )

    draw_word_cluster(left_tools, x0 + col0 + 18, col1 - 36, ["#8f5324", "#c7782d", "#7a4a1c", "#d99f5c"])
    draw_word_cluster(right_tools, x0 + col0 + col1 + 18, col2 - 36, ["#1f6f63", "#2a9d8f", "#335c67", "#5aa9a1"])
    save_svg(path, parts)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)


def add_table_from_df(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for col_idx, column in enumerate(df.columns):
        table.cell(0, col_idx).text = str(column)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = str(value)


def build_report(
    df: pd.DataFrame,
    correlations: pd.DataFrame,
    low_full: pd.DataFrame,
    high_not_full: pd.DataFrame,
    maturity_summary: pd.DataFrame,
    tool_summary: pd.DataFrame,
) -> str:
    adopted = df[df["is_adopted"]].copy()
    quantity_perf = float(correlations.loc["Number of Tools Adopted", "Performance Score"])
    quality_perf = float(correlations.loc["Maturity Score", "Performance Score"])
    quantity_quality = float(correlations.loc["Number of Tools Adopted", "Maturity Score"])
    integration_perf = float(correlations.loc["Integration Score", "Performance Score"])

    lines = [
        "Exploration of maturity, quantity, integration, and performance",
        "",
        f"Sample size: {len(df)} total rows, {len(adopted)} rows with Adoption Status = Adopted.",
        "",
        "Key row-level pattern:",
        f"- Low maturity but pure full integration appears in {len(low_full)} adopted cases.",
        f"- High maturity but still not pure full integration appears in {len(high_not_full)} adopted cases.",
        "",
        "Core correlations among adopted cases:",
        f"- Quantity (number of tools adopted) with quality (maturity score): r = {quantity_quality:.3f}",
        f"- Quantity with performance score: r = {quantity_perf:.3f}",
        f"- Quality with performance score: r = {quality_perf:.3f}",
        f"- Integration score with performance score: r = {integration_perf:.3f}",
        "",
        "Interpretation:",
        "Quantity and quality move together strongly, but they are not the same thing. Organizations with more tools usually also report more mature tools, yet the mismatched rows show that a small bundle of low-quality tools can still be reported as fully integrated, while a large high-quality portfolio can remain only partly integrated.",
        "This supports the regression interpretation that integration is an implementation process rather than a direct synonym for tool quality. When maturity is high, additional integration effort may bring coordination costs, interoperability burdens, or transition friction that dampen performance gains.",
        "For substantive interpretation, the number of tools adopted can be treated as QUANTITY, while technological maturity captures QUALITY. The raw data suggest that quantity alone is not enough. Higher performance is more consistent when institutions combine enough tools with better-quality tools and a realistic integration pathway.",
        "",
        "Most visible contradictory low-maturity but full-integration cases:",
    ]
    for _, row in low_full.iterrows():
        lines.append(
            f"- {row['Institution']}: maturity {row['Maturity Score']:.0f}, quantity {row['Number of Tools Adopted']:.0f}, integration level {row['System Integration Level']}, performance {row['Performance Score']:.1f}"
        )
    lines.extend(["", "Most visible high-maturity but not-pure-full cases:"])
    for _, row in high_not_full.iterrows():
        lines.append(
            f"- {row['Institution']}: maturity {row['Maturity Score']:.0f}, quantity {row['Number of Tools Adopted']:.0f}, integration level {row['System Integration Level']}, performance {row['Performance Score']:.1f}"
        )
    lines.extend(
        [
            "",
            "Tool-family reading:",
            "Common foundational tools such as GIS and AMI appear across low and moderate maturity settings, which helps explain why some low-maturity institutions still claim full integration. More advanced tools such as SCADA, IoT sensors, analytics, and predictive maintenance sit at higher average maturity, but they are often embedded in institutions that still report transitional integration states rather than simple full integration.",
            "",
            "Practical takeaway:",
            "The raw rows suggest a progression from quantity to quality to integration discipline. Adding more tools expands digital reach, better-quality tools deepen technical capability, and integration determines whether that capability is translated into stable performance.",
        ]
    )
    return "\n".join(lines) + "\n"


def build_document(report_text: str, low_full: pd.DataFrame, high_not_full: pd.DataFrame, tool_summary: pd.DataFrame) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Maturity, quantity, and integration interpretation")
    run.bold = True
    run.font.size = Pt(15)

    intro = doc.add_paragraph(
        "This note re-examines the row-level data in original_data.xlsx to understand why technological maturity and technology integration do not move together perfectly."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_heading(doc, "Main interpretation")
    for paragraph in report_text.split("\n\n")[:4]:
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_heading(doc, "Contradictory cases", level=2)
    doc.add_paragraph("Low maturity but full integration")
    add_table_from_df(doc, low_full.reset_index(drop=True))
    doc.add_paragraph("")
    doc.add_paragraph("High maturity but not pure full integration")
    add_table_from_df(doc, high_not_full.reset_index(drop=True))

    add_heading(doc, "Tool-family signal", level=2)
    add_table_from_df(doc, tool_summary.head(10).reset_index(drop=True))

    closing = doc.add_paragraph(
        "In short, number of tools adopted should be read as quantity, while maturity score captures quality. Integration then reflects how well those tools are assembled into working practice. The raw data imply that quality and integration should not be treated as interchangeable."
    )
    closing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(DOCX_FILE)


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    PLOTS_DIR.mkdir(exist_ok=True)

    df = load_data()
    adopted = df[df["is_adopted"]].copy()
    correlations = correlation_table(adopted)
    low_full, high_not_full = contradiction_tables(df)
    maturity_summary = maturity_integration_summary(df)
    perf_summary = performance_summary(df)
    tool_summary = tool_level_summary(df)

    df.to_csv(OUTPUT_DIR / "cleaned_original_rows.csv", index=False)
    correlations.to_csv(OUTPUT_DIR / "adopted_correlations.csv")
    low_full.to_csv(OUTPUT_DIR / "low_maturity_full_integration_cases.csv", index=False)
    high_not_full.to_csv(OUTPUT_DIR / "high_maturity_not_full_cases.csv", index=False)
    maturity_summary.to_csv(OUTPUT_DIR / "maturity_integration_summary.csv", index=False)
    perf_summary.to_csv(OUTPUT_DIR / "performance_by_maturity.csv", index=False)
    tool_summary.to_csv(OUTPUT_DIR / "tool_level_summary.csv", index=False)

    heatmap_svg(maturity_summary, PLOTS_DIR / "figure1_maturity_integration_heatmap.svg")
    scatter_svg(df, PLOTS_DIR / "figure2_quantity_performance_quality.svg")
    tool_map_svg(tool_summary, PLOTS_DIR / "figure3_tool_family_map.svg")
    contradiction_panel_svg(low_full, high_not_full, PLOTS_DIR / "figure4_contradictory_cases.svg")

    report_text = build_report(df, correlations, low_full, high_not_full, maturity_summary, tool_summary)
    REPORT_FILE.write_text(report_text, encoding="utf-8")
    build_document(report_text, low_full, high_not_full, tool_summary)


if __name__ == "__main__":
    main()
